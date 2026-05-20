from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "Pro_Gamer_Physical_Test_게임기획서.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(10.5)


def add_kv_table(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    for key, value in rows:
        cells = table.add_row().cells
        cells[0].width = Inches(1.55)
        cells[1].width = Inches(4.95)
        set_cell_shading(cells[0], "F2F2F2")
        set_cell_text(cells[0], key, True)
        set_cell_text(cells[1], value)
        cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.add_paragraph()


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(10.5)


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(10.5)


def add_para(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.1
    run = p.add_run(text)
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(10.5)


def add_heading(doc, text, level):
    p = doc.add_heading("", level=level)
    run = p.add_run(text)
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.color.rgb = RGBColor(46, 116, 181) if level <= 2 else RGBColor(31, 77, 120)
    run.font.bold = True
    if level == 1:
        run.font.size = Pt(16)
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(7)
    elif level == 2:
        run.font.size = Pt(13)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(5)
    else:
        run.font.size = Pt(12)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)


def main():
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("게임 기획서")
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(24)
    run.font.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(18)
    run = subtitle.add_run("Pro Gamer Physical Test")
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(85, 85, 85)

    add_heading(doc, "1. 게임 개요", 1)
    add_kv_table(
        doc,
        [
            ("게임 제목", "Pro Gamer Physical Test"),
            ("장르", "미니게임 / 반응속도 테스트 / 콘솔 게임"),
            ("플랫폼", "PC"),
            ("개발 언어", "C언어"),
            ("개발 환경", "Windows 콘솔 / Visual Studio 기준"),
            ("목표 유저", "게임 피지컬 테스트에 관심 있는 사람, 프로게이머 지망생 느낌으로 가볍게 테스트하고 싶은 사람"),
        ],
    )

    add_heading(doc, "2. 게임 콘셉트 및 핵심 특징", 1)
    add_para(doc, "여러 가지 간단한 피지컬 테스트를 진행하고, 최종 점수로 플레이어의 게임 티어를 판정하는 게임이다.")
    add_bullet(doc, "반응속도, 정확도, 순발력 등을 테스트한다.")
    add_bullet(doc, "각 미니게임 점수를 합산해서 최종 점수를 만든다.")
    add_bullet(doc, "결과는 롤 티어처럼 IRON, BRONZE, SILVER, GOLD, PLATINUM, DIAMOND 등으로 표시한다.")
    add_bullet(doc, "그냥 점수만 나오는 것보다 내가 어느 정도인지 바로 알 수 있어서 재미가 있다.")

    add_heading(doc, "3. 스토리 및 세계관", 1)
    add_para(doc, "스토리가 중요한 게임은 아니다. 플레이어는 프로게이머 선발 테스트에 참가한 사람이라는 설정이다.")
    add_para(doc, "여러 테스트를 통과하면서 본인이 프로게이머가 될 기본 피지컬이 있는지 확인한다. 분위기는 e스포츠 훈련장이나 테스트장 느낌으로 잡는다.")
    add_bullet(doc, "주인공: 테스트를 보는 플레이어")
    add_bullet(doc, "목표: 높은 점수를 받아 높은 티어를 받는 것")
    add_bullet(doc, "분위기: 게임 훈련장, e스포츠 테스트장")

    add_heading(doc, "4. 게임 플레이 및 시스템", 1)
    add_heading(doc, "4-1. 미니게임 구성", 2)
    add_kv_table(
        doc,
        [
            ("반응속도 테스트", "화면에 신호가 나오면 최대한 빨리 스페이스바를 누른다. 너무 빨리 누르면 실패 처리한다."),
            ("방향키 순발력 테스트", "화면에 랜덤 방향키가 나오면 해당 방향키를 빠르게 입력한다."),
            ("마우스 광클 테스트", "일정 시간 동안 마우스 왼쪽 버튼을 최대한 많이 클릭한다. FPS에서 말하는 광클 피지컬을 테스트한다."),
            ("타이밍 테스트", "표시가 랜덤한 시간에 위에서 아래로 한 번 떨어진다. 지정된 위치에 가깝게 멈출수록 높은 점수."),
        ],
    )

    add_heading(doc, "4-2. 조작 방식", 2)
    add_bullet(doc, "방향키: 이동 및 방향 입력")
    add_bullet(doc, "스페이스바: 선택 / 반응 입력")
    add_bullet(doc, "ESC: 종료")

    add_heading(doc, "4-3. 게임 진행 방식", 2)
    add_number(doc, "메인 메뉴 출력")
    add_number(doc, "테스트 시작")
    add_number(doc, "원하는 미니게임을 선택해서 플레이")
    add_number(doc, "4개의 미니게임을 모두 플레이")
    add_number(doc, "각 테스트별 점수 계산")
    add_number(doc, "총점 계산")
    add_number(doc, "4개 테스트가 모두 끝나면 총점 계산")
    add_number(doc, "최종 티어 출력")
    add_number(doc, "다시 하기 또는 종료 선택")

    add_heading(doc, "5. 그래픽 및 사운드 스타일", 1)
    add_bullet(doc, "아트 스타일: 콘솔 문자 기반 ASCII 그래픽")
    add_bullet(doc, "카메라 시점: 고정 화면")
    add_bullet(doc, "그래픽 방식: printf, gotoxy를 사용해서 화면에 문자 출력")
    add_bullet(doc, "사운드: 필수는 아니지만 시간이 남으면 Beep() 함수로 성공/실패 효과음 추가")

    add_heading(doc, "6. UI/UX 디자인", 1)
    add_para(doc, "UI는 복잡하게 만들지 않고 콘솔 화면에서 메뉴와 결과가 확실히 보이도록 만든다.")
    add_bullet(doc, "시작 화면")
    add_bullet(doc, "게임 설명")
    add_bullet(doc, "테스트 시작")
    add_bullet(doc, "결과 화면")
    add_bullet(doc, "종료")

    add_heading(doc, "결과 화면 예시", 2)
    result = doc.add_paragraph()
    result.paragraph_format.space_after = Pt(8)
    result_run = result.add_run(
        "===== RESULT =====\n\n"
        "Reaction Test : 85\n"
        "Arrow Test    : 72\n"
        "Click Test    : 78\n"
        "Timing Test   : 90\n\n"
        "Total Score   : 315\n"
        "Your Tier     : GOLD"
    )
    result_run.font.name = "Consolas"
    result_run.font.size = Pt(10)

    add_heading(doc, "티어 기준 예시", 2)
    add_kv_table(
        doc,
        [
            ("0 ~ 99", "IRON"),
            ("100 ~ 199", "BRONZE"),
            ("200 ~ 299", "SILVER"),
            ("300 ~ 399", "GOLD"),
            ("400 ~ 499", "PLATINUM"),
            ("500 이상", "DIAMOND"),
        ],
    )

    add_heading(doc, "7. 비즈니스 모델", 1)
    add_para(doc, "과제용 C언어 콘솔 게임이므로 비즈니스 모델은 따로 없다.")
    add_bullet(doc, "유료/무료 여부: 무료")
    add_bullet(doc, "수익 모델: 없음")
    add_bullet(doc, "목적: C언어 게임 개발 과제 제출")

    add_heading(doc, "8. 개발 일정 및 팀 구성", 1)
    add_kv_table(
        doc,
        [
            ("1주차", "기획서 작성, 게임 구조 설계"),
            ("2주차", "메인 메뉴, 화면 출력, 키 입력 구현"),
            ("3주차", "반응속도 테스트 구현"),
            ("4주차", "방향키 테스트, 마우스 광클 테스트 구현"),
            ("5주차", "타이밍 테스트 구현"),
            ("6주차", "점수 계산, 티어 출력, 버그 수정"),
        ],
    )
    add_para(doc, "팀 구성은 개인 개발로 진행한다. 기획, 프로그래밍, 테스트, 발표 자료 정리를 혼자 맡는다.")

    add_heading(doc, "9. 리스크 및 해결 방안", 1)
    add_heading(doc, "9-1. Mac과 Windows 환경 차이", 2)
    add_para(doc, "수업 예제 코드는 windows.h, conio.h, SetConsoleCursorPosition()을 사용한다. 이 함수들은 Windows 전용이라 Mac에서는 그대로 실행하기 어렵다.")
    add_para(doc, "최종 제출은 수업 환경에 맞춰 Windows 콘솔 기준으로 제작한다. Mac에서는 코드 작성과 구조 설계 위주로 작업하고, 실행 테스트는 Windows 또는 학교 컴퓨터의 Visual Studio에서 진행한다.")
    add_para(doc, "가능하면 gotoxy(), getKeyDown() 같은 입력/출력 관련 함수는 따로 분리해서 나중에 수정하기 쉽게 만든다.")

    add_heading(doc, "9-2. 미니게임 개수 문제", 2)
    add_para(doc, "처음부터 미니게임을 더 많이 넣으면 완성도가 떨어질 수 있다.")
    add_para(doc, "미니게임은 반응속도 테스트, 방향키 순발력 테스트, 마우스 광클 테스트, 타이밍 테스트 4개로 고정한다. 대신 플레이 순서는 사용자가 선택할 수 있게 하고, 4개를 모두 완료해야 최종 티어가 나오도록 만든다.")

    add_heading(doc, "10. 참고 게임 및 자료", 1)
    add_bullet(doc, "Human Benchmark")
    add_bullet(doc, "Aim Lab")

    doc.save(OUT)


if __name__ == "__main__":
    main()
